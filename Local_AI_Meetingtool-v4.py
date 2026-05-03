#文字おこし、話者分離、要約に対応したバージョン（要約はLM Studio使用）
#使用時はLM Studioでローカルサーバーを起動し、llamaを使用すること
#Windowsでは動作しないので注意
#M2チップ16GBに合わせた設定（メモリ最適化済み）
import customtkinter as ctk
from tkinter import filedialog, messagebox
import mlx_whisper
import threading
import os
import docx
import json
import re
import gc
import torch
import openai
import importlib
from datetime import datetime #自動保存用のタイムスタンプ取得

class CancelledError(Exception):
    pass

# --- 環境設定 ---
os.environ["HF_HUB_OFFLINE"] = "0"
# pyannoteは重いため、話者分離を使う場合だけ遅延ロードする

# --- デザインの基本設定 ---
ctk.set_appearance_mode("light") 
ctk.set_default_color_theme("blue")

class Whisperapp:
    SUMMARY_HEADER = "=========================\n【要約結果】\n========================="
    SUMMARY_SECTION_PATTERN = re.compile(r"\n\s*\n=+\n【要約結果】\n=+\n.*\Z", re.DOTALL)

    def __init__(self, root):
        self.root = root
        self.root.title("Local AI MeetingTool Ver.4 - 文字起こし & 要約")
        self.root.geometry("1000x800")
        self.root.configure(fg_color="#FFFFFF")
        self.is_closing = False

        # Mac用フォント設定
        self.font_title = ("Hiragino Sans", 15, "bold")
        self.font_main = ("Hiragino Sans", 13)
        self.font_text = ("Menlo", 13)

        self.diarization_pipeline = None
        self.processing_thread = None
        self.cancel_event = threading.Event()
        self.active_task = None
        self.keyring = self.load_keyring()
        self.keyring_service = "Local AI MeetingTool Ver.3"
        self.keyring_username = "huggingface_token"
        self.diarization_batch_size = 64 # M2チップ16GB向けの話者分離(pyannote)バッチサイズ
        self.context_length = 4000 # M2チップ16GB向けにLM Studioへ渡す1チャンクあたりの目安文字数
        self.summary_timeout_seconds = 1800 #要約する際のタイムアウト時間
        self.summary_chunk_chars = 4000 # 要約する文章を分割しLM Studioに渡す（メモリ節約のため小さく設定）
        self.summary_merge_chunk_chars = 3000 # 中間要約を統合する際の入力上限
        self.summary_merge_group_items = 5 # 中間要約を一度に統合する最大件数
        self.summary_intermediate_target_chars = 2500 # 階層統合中の中間要約目安
        self.summary_final_target_chars = 7000 # 最終要約の目安
        self.summary_intermediate_max_tokens = 2048
        self.summary_final_max_tokens = 4096
        self.summary_max_merge_levels = 8
        self.transcript_text = ""
        self.summary_text = ""


        self.models = {
           "Large v3（最高精度・低速）": "mlx-community/whisper-large-v3-mlx", 
           "Turbo（高速・高精度）": "mlx-community/whisper-large-v3-turbo",
            "Small（バランス）": "mlx-community/whisper-small-mlx",
            "Base（軽量・高速）": "mlx-community/whisper-base-mlx",
            "Tiny（最速・低精度）": "mlx-community/whisper-tiny-mlx"
        }

        self.summary_prompts = {
            "標準要約（事実のみ簡潔に）": "あなたは優秀なアシスタントです。以下の文字起こしテキストを要約してください。不要な相槌や重複した議論は削り、事実のみを正確にまとめてください。",
            "議事録（箇条書き）": "あなたは優秀な書記です。以下の文字起こしテキストから重要なポイントを抽出し、見出しと箇条書きを使った分かりやすい議事録を作成してください。",
            "決定事項・ToDoの抽出": "以下の文字起こしテキストから、「決定された事項」と「各担当者の次のアクション(ToDo)」のみを箇条書きで分かりやすく抽出してください。",
            "行政報告書形式（きわめてフォーマル）": "あなたは優秀な行政官です。以下の文字起こしテキストを、公式な行政報告書にふさわしい、厳格で正確な公用文にて要約してください。客観的な事実関係を整理し、感情表現や重複を完全に排除して簡潔にまとめてください。",
            "業務マニュアル・引き継ぎ書（Markdown）": """あなたは自治体の優秀な業務改善担当者です。以下の文字起こしテキストの内容を読み解き、後任者が読んですぐに作業を再現できるように、業務マニュアル・引き継ぎ書としてまとめてください。
            出力はMarkdown形式（README風）とし、適宜「#」などの見出し、箇条書き、手順のステップ分けを用いて、視覚的で読みやすく整理してください。
            出力には以下の項目を必ず含めてください：\n- 業務の目的・概要\n- 必要な前提知識・システム権限\n- 具体的な作業手順（ステップバイステップ）\n- 注意点・つまずきやすいポイント"""
        }
        self.summary_modes = ["分割要約（長文向け）", "一括要約（短文向け）"]

        # --- メインレイアウト ---
        self.main_frame = ctk.CTkFrame(root, fg_color="transparent")
        self.main_frame.pack(fill=ctk.BOTH, expand=True, padx=30, pady=15) 

        # 1. ファイル選択セクション
        self.label_step1 = ctk.CTkLabel(self.main_frame, text="1. 音声・動画ファイルを選択", font=self.font_title, text_color="#1D1D1F")
        self.label_step1.pack(anchor="w", pady=(0, 5))

        self.select_btn = ctk.CTkButton(self.main_frame, text="ファイルを選択", command=self.select_file, 
                                        font=self.font_main, height=35, corner_radius=8, 
                                        fg_color="#F5F5F7", text_color="#007AFF", hover_color="#E5E5E7")
        self.select_btn.pack(fill=ctk.X)

        self.file_path_label = ctk.CTkLabel(self.main_frame, text="ファイルが選択されていません", 
                                            text_color="#8E8E93", font=("Hiragino Sans", 12))
        self.file_path_label.pack(pady=(2, 10))

        # 2. 設定エリア
        self.configure_frame = ctk.CTkFrame(self.main_frame, fg_color="#FBFBFD", corner_radius=15, border_width=1, border_color="#D2D2D7")
        self.configure_frame.pack(fill=ctk.X, pady=5, ipady=5)

        # Hugging Face トークン
        ctk.CTkLabel(self.configure_frame, text="Hugging Face トークン", font=self.font_main).grid(row=0, column=0, padx=15, pady=(10, 5), sticky="w")
        self.token_entry = ctk.CTkEntry(self.configure_frame, width=350, placeholder_text="Token (ローカル認証済なら不要)",
                                         corner_radius=8, border_color="#D2D2D7", fg_color="#FFFFFF", show="*")
        self.token_entry.grid(row=0, column=1, padx=15, pady=(10, 5), sticky="w")

        # モデル選択
        ctk.CTkLabel(self.configure_frame, text="使用するAIモデル", font=self.font_main).grid(row=1, column=0, padx=15, pady=5, sticky="w")
        self.model_var = ctk.StringVar(value="Turbo（高速・高精度）")  # M2チップ16GB向けのデフォルト
        self.model_menu = ctk.CTkComboBox(self.configure_frame, values=list(self.models.keys()), 
                                          variable=self.model_var, width=350, corner_radius=8)
        self.model_menu.grid(row=1, column=1, padx=15, pady=5, sticky="w")

        # 話者分離オプション
        self.diarize_var = ctk.BooleanVar(value=False)
        self.diarize_check = ctk.CTkCheckBox(self.configure_frame, text="話者分離を有効にする", variable=self.diarize_var, font=self.font_main)
        self.diarize_check.grid(row=1, column=2, padx=(15, 5), pady=5, sticky="w")

        self.num_speakers_var = ctk.StringVar(value="自動")
        self.num_speakers_menu = ctk.CTkComboBox(self.configure_frame, values=["自動", "2", "3", "4", "5", "6", "7", "8"], variable=self.num_speakers_var, width=100)
        self.num_speakers_menu.grid(row=1, column=3, padx=(5, 15), pady=5, sticky="w")

        # 3. 専門用語辞書
        self.dict_frame = ctk.CTkFrame(self.main_frame, fg_color="transparent")
        self.dict_frame.pack(fill=ctk.X, pady=(10, 5))

        ctk.CTkLabel(self.dict_frame, text="3. 専門用語・略語の登録（任意）", 
                     font=self.font_title, text_color="#1D1D1F").pack(side=ctk.LEFT)

        #専門用語辞書の読み込みと保存
        self.save_dict_btn = ctk.CTkButton(self.dict_frame, text="保存", command=self.save_dictionary, width=60, height=28, font=self.font_main, fg_color="#F5F5F7", text_color="#007AFF", hover_color="#E5E5E7")
        self.save_dict_btn.pack(side=ctk.RIGHT, padx=(5, 0))

        self.load_dict_btn = ctk.CTkButton(self.dict_frame, text="読込", command=self.load_dictionary, width=60, height=28, font=self.font_main, fg_color="#F5F5F7", text_color="#007AFF", hover_color="#E5E5E7")
        self.load_dict_btn.pack(side=ctk.RIGHT)

        self.prompt_entry = ctk.CTkEntry(self.main_frame, placeholder_text="例）AI: 人工知能, ML: 機械学習", 
                                         corner_radius=8, height=35, border_color="#D2D2D7", fg_color="#FFFFFF")
        self.prompt_entry.pack(fill=ctk.X, pady=(0, 10))

        # 4. 発言者名の置換
        self.speaker_frame = ctk.CTkFrame(self.main_frame, fg_color="transparent")
        self.speaker_frame.pack(fill=ctk.X, pady=(0, 5))

        ctk.CTkLabel(self.speaker_frame, text="4. 発言者名の置換（任意）",
                     font=self.font_title, text_color="#1D1D1F").pack(side=ctk.LEFT)

        self.apply_speaker_btn = ctk.CTkButton(
            self.speaker_frame,
            text="置換実行",
            command=self.apply_speaker_names_from_ui,
            state="disabled",
            width=80,
            height=28,
            font=self.font_main,
            fg_color="#F5F5F7",
            text_color="#007AFF",
            hover_color="#E5E5E7",
        )
        self.apply_speaker_btn.pack(side=ctk.RIGHT, padx=(5, 0))

        self.extract_speaker_btn = ctk.CTkButton(
            self.speaker_frame,
            text="候補抽出",
            command=self.extract_speaker_names_from_result,
            state="disabled",
            width=80,
            height=28,
            font=self.font_main,
            fg_color="#F5F5F7",
            text_color="#007AFF",
            hover_color="#E5E5E7",
        )
        self.extract_speaker_btn.pack(side=ctk.RIGHT)

        self.speaker_entry = ctk.CTkEntry(
            self.main_frame,
            placeholder_text="例）SPEAKER_00: 市民課長, SPEAKER_01: 委託事業者A, Unknown: 不明",
            corner_radius=8,
            height=35,
            border_color="#D2D2D7",
            fg_color="#FFFFFF",
        )
        self.speaker_entry.pack(fill=ctk.X, pady=(0, 10))

        # 実行・キャンセルボタン群
        self.run_frame = ctk.CTkFrame(self.main_frame, fg_color="transparent")
        self.run_frame.pack(fill=ctk.X, pady=5)

        self.run_btn = ctk.CTkButton(self.run_frame, text="文字起こしを開始する", command=self.start_thread, state="disabled",
                                     height=45, corner_radius=22, font=("Hiragino Sans", 15, "bold"), fg_color="#007AFF", hover_color="#005BB5")
        self.run_btn.pack(side=ctk.LEFT, expand=True, fill=ctk.X, padx=(0, 5))

        self.cancel_btn = ctk.CTkButton(self.run_frame, text="中止", command=self.cancel_process, state="disabled", width=80,
                                     height=45, corner_radius=22, font=("Hiragino Sans", 15, "bold"), text_color="#FFFFFF", fg_color="#FF3B30", hover_color="#D70015")
        self.cancel_btn.pack(side=ctk.LEFT, padx=(5, 0))

        # 結果表示エリア
        self.result_area = ctk.CTkTextbox(self.main_frame, height=150, corner_radius=10, border_width=1, 
                                          border_color="#D2D2D7", font=self.font_text, fg_color="#FFFFFF", text_color="#1D1D1F")
        self.result_area.pack(fill=ctk.BOTH, expand=True, pady=10)

        # 保存・要約ボタン群
        self.action_frame = ctk.CTkFrame(self.main_frame, fg_color="transparent")
        self.action_frame.pack(fill=ctk.X, pady=(0, 5))

        self.save_txt_btn = ctk.CTkButton(self.action_frame, text="Text保存", command=self.save_text, state="disabled", corner_radius=8, fg_color="#8E8E93")
        self.save_txt_btn.pack(side=ctk.LEFT, padx=5, expand=True, fill=ctk.X)

        self.save_word_btn = ctk.CTkButton(self.action_frame, text="Word保存", command=self.save_word, state="disabled", corner_radius=8, fg_color="#8E8E93")
        self.save_word_btn.pack(side=ctk.LEFT, padx=5, expand=True, fill=ctk.X)

        # 要約プロンプトの選択
        self.summary_prompt_var = ctk.StringVar(value="標準要約（事実のみ簡潔に）")
        self.summary_prompt_menu = ctk.CTkComboBox(self.action_frame, values=list(self.summary_prompts.keys()), 
                                                   variable=self.summary_prompt_var, corner_radius=8, width=280)
        self.summary_prompt_menu.pack(side=ctk.LEFT, padx=(15, 5))

        self.summary_mode_var = ctk.StringVar(value="分割要約（長文向け）")
        self.summary_mode_menu = ctk.CTkComboBox(self.action_frame, values=self.summary_modes, variable=self.summary_mode_var, corner_radius=8, width=160)
        self.summary_mode_menu.pack(side=ctk.LEFT, padx=(0, 5))

        self.summarize_btn = ctk.CTkButton(self.action_frame, text="LM Studioで要約", command=self.start_summarize_thread, state="disabled", corner_radius=8, fg_color="#34C759", hover_color="#248A3D")
        self.summarize_btn.pack(side=ctk.LEFT, padx=(0, 5), expand=True, fill=ctk.X)

        # ステータスバー（一番下に固定）
        self.status_label = ctk.CTkLabel(root, text="準備完了", fg_color="#F5F5F7", height=25, font=("Hiragino Sans", 11))
        self.status_label.pack(side=ctk.BOTTOM, fill=ctk.X)
        self.progress = ctk.CTkProgressBar(root, height=4, corner_radius=0, fg_color="#E5E5E7", progress_color="#007AFF")
        self.progress.pack(side=ctk.BOTTOM, fill=ctk.X)
        self.progress.set(0)

        self.filepath = ""
        self.config_filepath = "whisper_config.json"
        
        self.load_config()
        self.root.protocol("WM_DELETE_WINDOW", self.on_closing)

    #コンフィグの保存
    def save_config(self):
        self.save_hf_token()
        config = self.read_config_file()
        # 古い設定ファイルに残っている可能性があるトークンは、再保存時にJSONから除外する
        config.pop("hf_token", None)
        config = {
            **config,
            "model": self.model_var.get(),
            "diarize": self.diarize_var.get(),
            "num_speakers": self.num_speakers_var.get(),
            "prompt": self.prompt_entry.get(),
            "speaker_names": self.speaker_entry.get(),
            "summary_prompt": self.summary_prompt_var.get(),
            "summary_mode": self.summary_mode_var.get(),
            # 処理性能に関わる値は、画面項目ではなくwhisper_config.jsonから調整する
            "batch_size": self.diarization_batch_size,
            "context_length": self.context_length,
        }
        try:
            with open(self.config_filepath, "w", encoding="utf-8") as f:
                json.dump(config, f, ensure_ascii=False, indent=4)
        except Exception:
            pass

    def get_hf_token(self):
        if self.keyring is None:
            return None
        try:
            return self.keyring.get_password(self.keyring_service, self.keyring_username)
        except Exception:
            return None

    def load_keyring(self):
        try:
            return importlib.import_module("keyring")
        except ImportError:
            return None

    def load_diarization_pipeline_class(self):
        # 話者分離OFFの通常利用ではpyannoteをロードしない
        try:
            pyannote_audio = importlib.import_module("pyannote.audio")
            return pyannote_audio.Pipeline
        except ImportError as e:
            raise RuntimeError(
                "話者分離を使うには pyannote-audio が必要です。"
                "requirements.txt を使って依存関係をインストールしてください。"
            ) from e

    def save_hf_token(self):
        token = self.token_entry.get().strip()
        if self.keyring is None:
            if token:
                messagebox.showwarning(
                    "トークンを保存できません",
                    "keyring パッケージが見つからないため、Hugging Faceトークンは安全に保存されません。\n"
                    "requirements.txt を使って依存関係をインストールしてください。"
                )
            return False

        try:
            if token:
                self.keyring.set_password(self.keyring_service, self.keyring_username, token)
            else:
                try:
                    self.keyring.delete_password(self.keyring_service, self.keyring_username)
                except self.keyring.errors.PasswordDeleteError:
                    pass
            return True
        except Exception as e:
            messagebox.showwarning("トークン保存エラー", f"Hugging Faceトークンを安全に保存できませんでした:\n{e}")
            return False

    #コンフィグの読み込み
    def read_config_file(self):
        # 設定ファイルが壊れていてもアプリ起動を止めず、デフォルト値で続行する
        if not os.path.exists(self.config_filepath):
            return {}
        try:
            with open(self.config_filepath, "r", encoding="utf-8") as f:
                config = json.load(f)
            return config if isinstance(config, dict) else {}
        except Exception:
            return {}

    def get_int_config(self, config, keys, default, min_value=1, max_value=None):
        # 旧キー名も受け付けられるよう、候補キーを順番に見て最初の有効な整数を採用する
        for key in keys:
            if key not in config:
                continue
            raw_value = config.get(key)
            try:
                if isinstance(raw_value, bool):
                    continue
                value = int(raw_value)
            except (TypeError, ValueError):
                continue

            if value < min_value:
                continue
            if max_value is not None:
                value = min(value, max_value)
            return value
        return default

    def load_runtime_config(self, config=None):
        # バッチサイズやコンテキスト長など、処理開始直前にも反映したい設定だけを読み込む
        config = config if config is not None else self.read_config_file()

        # batch_size: pyannoteの話者分離バッチサイズ。大きいほど速いがメモリを使う
        self.diarization_batch_size = self.get_int_config(
            config,
            ("batch_size", "diarization_batch_size"),
            self.diarization_batch_size,
            min_value=1,
            max_value=64,
        )

        # context_length: 要約時にLM Studioへ渡す1チャンクあたりの目安文字数
        self.context_length = self.get_int_config(
            config,
            ("context_length", "summary_context_length", "summary_chunk_chars"),
            self.context_length,
            min_value=1000,
            max_value=4000,
        )
        self.summary_chunk_chars = self.context_length
        # 統合要約では入力に余白を残すため、通常チャンクより少し小さめにする
        self.summary_merge_chunk_chars = max(1000, int(self.context_length * 0.75))

    def load_config(self):
        config = self.read_config_file()
        self.load_runtime_config(config)

        try:
            if "model" in config and config["model"] in self.models:
                self.model_var.set(config["model"])
            if "diarize" in config:
                self.diarize_var.set(config["diarize"])
            if "num_speakers" in config:
                self.num_speakers_var.set(config["num_speakers"])
            if "prompt" in config and config["prompt"]:
                self.prompt_entry.insert(0, config["prompt"])
            if "speaker_names" in config and config["speaker_names"]:
                self.speaker_entry.insert(0, config["speaker_names"])
            if "summary_prompt" in config and config["summary_prompt"] in self.summary_prompts:
                self.summary_prompt_var.set(config["summary_prompt"])
            if "summary_mode" in config and config["summary_mode"] in self.summary_modes:
                self.summary_mode_var.set(config["summary_mode"])
        except Exception:
            pass

        stored_token = self.get_hf_token()
        legacy_token = config.get("hf_token")
        if stored_token:
            self.token_entry.insert(0, stored_token)
        elif legacy_token:
            self.token_entry.insert(0, legacy_token)
            if self.save_hf_token():
                self.save_config()

    def on_closing(self):
        self.is_closing = True
        self.cancel_event.set()
        self.save_config()
        self.root.destroy()

    def safe_after(self, callback):
        if self.is_closing:
            return
        try:
            self.root.after(0,callback)
        except Exception:
            pass

    def has_result_text(self):
        return bool(self.result_area.get("1.0", ctk.END).strip())

    def update_action_buttons(self):
        state = "normal" if self.has_result_text() else "disabled"
        self.save_txt_btn.configure(state=state)
        self.save_word_btn.configure(state=state)
        self.summarize_btn.configure(state=state)
        self.extract_speaker_btn.configure(state=state)
        self.apply_speaker_btn.configure(state=state)

    def reset_ui_after_task(self, status_text, progress_value=0, keep_status=False):
        if self.progress.cget("mode") == "indeterminate":
            self.progress.stop()
            self.progress.configure(mode="determinate")

        self.progress.set(progress_value)
        if not keep_status:
            self.status_label.configure(text=status_text)
            
        # ボタン群の有効化
        self.select_btn.configure(state="normal")
        self.run_btn.configure(state="normal" if self.filepath else "disabled")
        self.summary_mode_menu.configure(state="normal")
        self.summary_prompt_menu.configure(state="normal")
        self.update_action_buttons()
        self.cancel_btn.configure(state="disabled")
        self.active_task = None

    def ensure_not_cancelled(self):
        if self.cancel_event.is_set():
            raise CancelledError()

    def clean_repeated_text(self, text):
        # Whisperがまれに起こす同一単語・短いフレーズの異常反復だけを控えめに圧縮する
        text = re.sub(r'([^\s\n])\1{9,}', r'\1', text)
        text = re.sub(r'([;:.])\1{5,}', r'\1', text)

        word_chars = r'ぁ-んァ-ヶ一-龥A-Za-z0-9ー'
        word_repeat_pattern = rf'([{word_chars}]{{2,20}})(?:[ 　、。,.，．]*\1){{2,}}'
        text = re.sub(word_repeat_pattern, r'\1', text)

        phrase_repeat_pattern = r'([^\n]{4,40}?)(?:\1){2,}'
        text = re.sub(phrase_repeat_pattern, r'\1', text)

        return text

    def strip_summary_section(self, text):
        return self.SUMMARY_SECTION_PATTERN.sub("", text or "").strip()

    def get_text_for_summary(self):
        current_text = self.result_area.get("1.0", ctk.END).strip()
        transcript_text = self.strip_summary_section(current_text)
        if transcript_text:
            self.transcript_text = transcript_text
            return transcript_text
        return self.transcript_text.strip()

    def format_summary_section(self, summary):
        return f"\n\n{self.SUMMARY_HEADER}\n{summary.strip()}\n"

    def get_summary_from_display_text(self, text):
        match = self.SUMMARY_SECTION_PATTERN.search(text or "")
        if not match:
            return ""

        summary_section = match.group(0)
        header_index = summary_section.find(self.SUMMARY_HEADER)
        if header_index == -1:
            return ""
        return summary_section[header_index + len(self.SUMMARY_HEADER):].strip()

    # 文字起こしが完了したらテキストファイルを自動保存
    def auto_save_text(self, text):
        if not self.filepath or not text.strip():
            return None

        base_dir = os.path.dirname(self.filepath)
        base_name = os.path.splitext(os.path.basename(self.filepath))[0]
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        save_path = os.path.join(base_dir, f"{base_name}_transcript_{timestamp}.txt")

        with open(save_path, "w", encoding="utf-8") as f:
            f.write(text)

        return save_path

    # 要約を行う文章を分割
    def split_text_for_summary(self, text, max_chars=None):
        max_chars = max_chars or self.summary_chunk_chars
        text = (text or "").strip()
        if not text:
            return []
        if len(text) <= max_chars:
            return [text]

        lines = text.splitlines()
        chunks = []
        current_lines = []
        current_length = 0

        for line in lines:
            line_length = len(line) + 1
            if current_lines and current_length + line_length > max_chars:
                chunks.append("\n".join(current_lines).strip())
                current_lines = []
                current_length = 0

            if line_length > max_chars:
                for start in range(0, len(line), max_chars):
                    part = line[start:start + max_chars].strip()
                    if part:
                        chunks.append(part)
                continue

            current_lines.append(line)
            current_length += line_length

        if current_lines:
            chunks.append("\n".join(current_lines).strip())

        return [chunk for chunk in chunks if chunk]

    # ユーザー辞書の設定
    def parse_dictionary_text(self, dictionary_text, allow_csv_pairs=False):
        replacements = {}
        hint_words = []
        invalid_entries = []

        def add_mapping(wrong_text, correct_text=None, raw_entry=""):
            wrong = wrong_text.strip()
            correct = wrong if correct_text is None else correct_text.strip()
            if not wrong or not correct: # 空の置換もとは事故につながるため明示的に弾く
                entry = raw_entry.strip()
                if entry:
                    invalid_entries.append(entry)
                return

            if wrong not in replacements:
                hint_words.append(wrong)
            replacements[wrong] = correct

        def split_mapping(entry):
            for separator in ("=>", "->", "\t", ":", "："):
                if separator in entry:
                    wrong, correct = entry.split(separator, 1)
                    return wrong, correct
            return None, None

        for line in dictionary_text.splitlines():
            line = line.strip()
            if not line:
                continue

            # 明示的な対応表がある行は、カンマ等で複数ペアを書けるようにする
            has_mapping_separator = any(separator in line for separator in ("=>", "->", "\t", ":", "："))
            if has_mapping_separator:
                entries = re.split(r"[、,，]", line)
                for entry in entries:
                    entry = entry.strip()
                    if not entry:
                        continue
                    wrong, correct = split_mapping(entry)
                    if wrong is None:
                        add_mapping(entry)
                    else:
                        add_mapping(wrong, correct, entry)
                continue

            entries = [entry.strip() for entry in re.split(r"[、,，]", line) if entry.strip()]
            # 通常入力の「AI、ML」は2つのヒント語として扱い、CSV読み込み時だけ2列を置換ペアにする
            if allow_csv_pairs and len(entries) == 2:
                add_mapping(entries[0], entries[1], line)
            else:
                for entry in entries:
                    add_mapping(entry)

        if invalid_entries:
            preview = " / ".join(invalid_entries[:3])
            raise ValueError(f"辞書入力に空の置換元または置換先があります: {preview}")

        # 「AI」と『AI技術」のような重なりは、長い語を先にマッチさせる
        sorted_replacements = dict(
            sorted(replacements.items(), key=lambda item: len(item[0]), reverse=True)
        )
        return sorted_replacements, hint_words

    def build_initial_prompt(self, hint_words, max_chars=1200):
        # Whisperの初期プロンプトが長くなりすぎないように先頭から一定量だけ渡す
        selected_words = []
        current_length = 0

        for word in hint_words:
            addition = len(word) + (2 if selected_words else 0)
            if current_length + addition > max_chars:
                break
            selected_words.append(word)
            current_length += addition

        return ", ".join(selected_words) if selected_words else None

    def build_dictionary_replacer(self, replacements):
        if not replacements:
            return lambda text: text
        # セグメントごとに正規表現を組み立て直さないように置換関数を一度だけ作る
        pattern = re.compile("|".join(re.escape(word) for word in replacements))
        return lambda text: pattern.sub(lambda match: replacements[match.group(0)], text or "")

    def format_dictionary_text(self, replacements, separator=", "):
        # 読込・保存時の表記揺れをアプリないの標準形式に揃える
        entries = []
        for wrong, correct in replacements.items():
            if wrong == correct:
                entries.append(wrong)
            else:
                entries.append(f"{wrong}: {correct}")
        return separator.join(entries)

    def parse_speaker_name_text(self, speaker_text):
        mappings = {}
        invalid_entries = []

        def split_mapping(entry):
            for separator in ("=>", "->", "\t", ":", "：", "="):
                if separator in entry:
                    source, target = entry.split(separator, 1)
                    return source, target
            return None, None

        for line in speaker_text.splitlines():
            line = line.strip()
            if not line:
                continue

            for entry in re.split(r"[、,，]", line):
                entry = entry.strip()
                if not entry:
                    continue

                source, target = split_mapping(entry)
                if source is None:
                    invalid_entries.append(entry)
                    continue

                source = source.strip().strip("【】")
                target = target.strip().strip("【】")
                if not source or not target:
                    invalid_entries.append(entry)
                    continue

                mappings[source] = target

        if invalid_entries:
            preview = " / ".join(invalid_entries[:3])
            raise ValueError(f"発言者名の対応表に不正な項目があります: {preview}")

        return dict(sorted(mappings.items(), key=lambda item: len(item[0]), reverse=True))

    def format_speaker_name_text(self, mappings, separator=", "):
        return separator.join(f"{source}: {target}" for source, target in mappings.items())

    def extract_speaker_labels(self, text):
        labels = []
        seen = set()
        label_pattern = re.compile(r"^【([^】\n]+)】\s*\[\d{2}:\d{2}(?::\d{2})?\]", re.MULTILINE)

        for match in label_pattern.finditer(text or ""):
            label = match.group(1).strip()
            if label and label not in seen:
                labels.append(label)
                seen.add(label)

        return labels

    def apply_speaker_name_mappings(self, text, mappings):
        if not text or not mappings:
            return text, 0

        updated_count = 0

        def replace_structured_label(match):
            nonlocal updated_count
            label = match.group(1).strip()
            replacement = mappings.get(label)
            if replacement and replacement != label:
                updated_count += 1
                return f"【{replacement}】"
            return match.group(0)

        updated_text = re.sub(
            r"【([^】\n]+)】(?=\s*\[\d{2}:\d{2}(?::\d{2})?\])",
            replace_structured_label,
            text,
        )

        generated_label_mappings = {
            source: target
            for source, target in mappings.items()
            if source != target and re.match(r"^(SPEAKER_\d+|Unknown)$", source)
        }
        if generated_label_mappings:
            pattern = re.compile(
                r"(?<![A-Za-z0-9_])("
                + "|".join(re.escape(source) for source in generated_label_mappings)
                + r")(?![A-Za-z0-9_])"
            )
            updated_text, bare_count = pattern.subn(
                lambda match: generated_label_mappings[match.group(1)],
                updated_text,
            )
            updated_count += bare_count

        return updated_text, updated_count

    def sync_text_state_from_display(self):
        display_text = self.result_area.get("1.0", ctk.END).strip()
        self.transcript_text = self.strip_summary_section(display_text)
        self.summary_text = self.get_summary_from_display_text(display_text)

    def extract_speaker_names_from_result(self):
        transcript_text = self.strip_summary_section(self.result_area.get("1.0", ctk.END))
        labels = self.extract_speaker_labels(transcript_text)
        if not labels:
            messagebox.showwarning("確認", "発言者ラベルを抽出できませんでした。話者分離済みの結果を表示してから実行してください。")
            return

        try:
            existing_mappings = self.parse_speaker_name_text(self.speaker_entry.get().strip())
        except Exception as e:
            messagebox.showerror("エラー", str(e))
            return

        ordered_mappings = {}
        for label in labels:
            ordered_mappings[label] = existing_mappings.get(label, label)
        for source, target in existing_mappings.items():
            if source not in ordered_mappings:
                ordered_mappings[source] = target

        self.speaker_entry.delete(0, ctk.END)
        self.speaker_entry.insert(0, self.format_speaker_name_text(ordered_mappings))
        self.save_config()
        messagebox.showinfo("候補抽出", "発言者ラベルを抽出しました。対応する名前を編集してから置換実行してください。")

    def apply_speaker_names_from_ui(self):
        text = self.result_area.get("1.0", ctk.END).strip()
        if not text:
            messagebox.showwarning("警告", "置換するテキストがありません。")
            return

        try:
            mappings = self.parse_speaker_name_text(self.speaker_entry.get().strip())
        except Exception as e:
            messagebox.showerror("エラー", str(e))
            return

        effective_mappings = {source: target for source, target in mappings.items() if source != target}
        if not effective_mappings:
            messagebox.showwarning("確認", "置換する発言者名の対応表がありません。")
            return

        updated_text, count = self.apply_speaker_name_mappings(text, effective_mappings)
        if updated_text == text:
            messagebox.showinfo("確認", "一致する発言者ラベルが見つかりませんでした。")
            return

        self.result_area.delete("1.0", ctk.END)
        self.result_area.insert(ctk.END, updated_text)
        self.sync_text_state_from_display()
        self.save_config()
        self.update_action_buttons()
        self.status_label.configure(text=f"発言者名を置換しました ({count}件)")
        messagebox.showinfo("成功", f"発言者名を置換しました。({count}件)")

    #実行ファイルフェーズ
    def select_file(self):
        file_types = [("Audio/Video files", "*.mp3 *.wav *.flac *.m4a *.mp4 *.mov *.mkv"), ("All files", "*.*")]
        selected_path = filedialog.askopenfilename(filetypes=file_types)
        if selected_path:
            self.filepath = selected_path
            self.file_path_label.configure(text=os.path.basename(self.filepath), text_color="black")
            self.run_btn.configure(state="normal")

    #処理中はボタンを非アクティブ化
    def start_thread(self):
        self.cancel_event = threading.Event()
        self.active_task = "transcribe"
        self.transcript_text = ""
        self.result_area.delete("1.0", ctk.END)
        self.select_btn.configure(state="disabled")
        self.run_btn.configure(state="disabled")
        self.save_txt_btn.configure(state="disabled") 
        self.save_word_btn.configure(state="disabled")
        self.extract_speaker_btn.configure(state="disabled")
        self.apply_speaker_btn.configure(state="disabled")
        self.summary_mode_menu.configure(state="disabled")
        self.summary_prompt_menu.configure(state="disabled")
        self.summarize_btn.configure(state="disabled")
        self.cancel_btn.configure(state="normal")
        self.status_label.configure(text="処理しています…")
        self.progress.set(0)
        # アプリ起動後にJSONを書き換えた場合も、次の文字起こし実行で反映する
        self.load_runtime_config()
        self.save_config()

        task_config = {
            "filepath": self.filepath,
            "token": self.token_entry.get().strip(),
            "model_label": self.model_var.get(),
            "do_diarize": self.diarize_var.get(),
            "num_speakers": self.num_speakers_var.get(),
            "user_prompt": self.prompt_entry.get().strip(),
            "speaker_names": self.speaker_entry.get().strip(),
            "batch_size": self.diarization_batch_size,
        }

        self.processing_thread = threading.Thread(target=self.run_process, args=(task_config,))
        self.processing_thread.daemon = True 
        self.processing_thread.start()

    #処理の中止
    def cancel_process(self):
        if hasattr(self, 'processing_thread') and self.processing_thread and self.processing_thread.is_alive():
            self.status_label.configure(text="中止しています...")
            self.cancel_btn.configure(state="disabled")
            self.cancel_event.set()

    #実行処理フェーズ
    def run_process(self, task_config):
        try:
            filepath = task_config["filepath"]
            token = task_config["token"]
            do_diarize = task_config["do_diarize"]
            user_prompt = task_config["user_prompt"]
            speaker_name_mappings = self.parse_speaker_name_text(task_config["speaker_names"])
            
            if not filepath or not os.path.isfile(filepath):
                raise ValueError("ファイルが選択されていないか、存在しません。もう一度ファイルを選択し直してください。")
                
            self.ensure_not_cancelled()

            prompt_dict, hint_words = self.parse_dictionary_text(user_prompt)
            initial_prompt_str = self.build_initial_prompt(hint_words)

            # 1.Whisperで文字起こし
            self.safe_after(lambda: self.status_label.configure(text="Whisperで文字起こし中... しばらくお待ちください..."))
            self.safe_after(lambda: self.progress.configure(mode="indeterminate")) #プログレスバーのアニメーション
            self.safe_after(lambda: self.progress.start())
            model_path = self.models[task_config["model_label"]]

            #長時間データの場合の安定化オプション（M2チップ16GB向けにメモリ節約設定）
            whisper_options = {
                "language": "ja",                       #言語を日本語に固定
                "condition_on_previous_text": False,   #前の文脈を引きずらない
                "compression_ratio_threshold": 1.8,     #読み間違いのループを検知してリセット
                "no_speech_threshold": 0.7,             #無音区間での暴走を防ぐ
                "logprob_threshold": -0.8,              #無音区間の暴走抑止
                "temperature": 0.0,                      #ランダムな記号の生成抑止
                "word_timestamps": False,               #タイムスタンプ計算のバグによる記号生成の抑止
                "best_of": 1                            #メモリ節約のため候補数を1に設定
            }
            # Whisperで文字起こし
            whisper_result = mlx_whisper.transcribe(
                filepath,
                path_or_hf_repo=model_path,
                initial_prompt=initial_prompt_str or None,
                **whisper_options,
            )
            
            # 途中でキャンセルされていないか確認
            self.ensure_not_cancelled()
            self.safe_after(lambda: self.progress.stop())
            self.safe_after(lambda: self.progress.configure(mode="determinate"))
            self.safe_after(lambda: self.progress.set(0))

            #専門用語の置換
            if prompt_dict:
                replace_dictionary_terms = self.build_dictionary_replacer(prompt_dict)
                for segment in whisper_result.get("segments", []):
                    segment["text"] = replace_dictionary_terms(segment.get("text", ""))

                whisper_result["text"] = replace_dictionary_terms(whisper_result.get("text", ""))

            final_text = ""
            gc.collect() 
            if torch.backends.mps.is_available():
                torch.mps.empty_cache()  # M2チップ向けにメモリクリアを強化

            # 2.話者分離の実行
            if do_diarize:
                self.ensure_not_cancelled() # キャンセルされていないか確認
                self.safe_after(lambda: self.status_label.configure(text="話者解析中... (初回はモデルをロードします)"))
                if self.diarization_pipeline is None:
                    auth_param = token if token else None
                    try:
                        Pipeline = self.load_diarization_pipeline_class()
                        self.diarization_pipeline = Pipeline.from_pretrained("pyannote/speaker-diarization-3.1", token=auth_param)

                        if torch.backends.mps.is_available() and hasattr(self.diarization_pipeline, "to"):
                            self.diarization_pipeline.to(torch.device("mps"))
                    except Exception as e:
                        raise RuntimeError(f"話者分離モデルの読み込みに失敗しました:\n{e}")
                if hasattr(self.diarization_pipeline, "set_batch_size"):
                    # 既にモデルを読み込み済みの場合でも、最新のJSON設定を毎回反映する
                    self.diarization_pipeline.set_batch_size(task_config["batch_size"])

                def diarization_hook(*args, **kwargs):
                    if self.cancel_event.is_set():
                        raise CancelledError()
                    completed = kwargs.get("completed") or 0
                    total = kwargs.get("total") or 1

                    #話者分離のプログレスバーアニメーション
                    if total >0:
                        percentage_float = completed / total
                        percentage_int = int(percentage_float * 100)

                        self.safe_after(lambda: self.progress.set(percentage_float))
                        self.safe_after(lambda: self.status_label.configure(
                            text=f"話者解析中...しばらくお待ちください... ({percentage_int}%)"))

                num_spk = None if task_config["num_speakers"] == "自動" else int(task_config["num_speakers"])
                diarization_result = self.diarization_pipeline(filepath, num_speakers=num_spk, hook=diarization_hook)

                self.ensure_not_cancelled()
                
                #話者分離のトラックを抽出
                tracks = [] 
                if hasattr(diarization_result, "itertracks"): 
                    for turn, _, spk_label in diarization_result.itertracks(yield_label=True):
                        tracks.append((turn, spk_label))
                elif hasattr(diarization_result, "speaker_diarization"): 
                    data = diarization_result.speaker_diarization
                    if hasattr(data, "itertracks"):
                        for turn, _, spk_label in data.itertracks(yield_label=True):
                            tracks.append((turn, spk_label))
                    else:
                        for turn, spk_label in data:
                            tracks.append((turn, spk_label))
                else:
                    raise ValueError("Pyannoteの結果の形式が不明です。")

                # セグメントごとに話者を割り当て
                for segment in whisper_result["segments"]:
                    start_time = segment["start"]
                    end_time = segment["end"]
                    text = segment["text"].strip()
                    speaker = "Unknown"
                    max_overlap = -1.0

                    for turn, spk_label in tracks:
                        overlap_start = max(start_time, turn.start) 
                        overlap_end = min(end_time, turn.end)
                        overlap = overlap_end - overlap_start
                        if overlap > max_overlap and overlap > 0: 
                            max_overlap = overlap
                            speaker = spk_label

                    time_str = f"[{int(start_time//60):02d}:{int(start_time%60):02d}]"
                    display_speaker = speaker_name_mappings.get(speaker, speaker)
                    final_text += f"【{display_speaker}】{time_str} {text}\n"
            
            # 話者分離を行わない場合はWhisper結果をそのまま返す
            else:
                final_text = whisper_result["text"].strip()

            final_text = self.clean_repeated_text(final_text)
            gc.collect()  # メモリクリアを追加
            if torch.backends.mps.is_available():
                torch.mps.empty_cache()

            self.ensure_not_cancelled() #最終結果表示前にキャンセルされていないか確認
            self.safe_after(lambda: self.show_result(final_text))

        except CancelledError:
            gc.collect()
            if torch.backends.mps.is_available(): torch.mps.empty_cache()
            self.safe_after(lambda: self.reset_ui_after_task("処理が中止されました"))

        #エラー表示
        except Exception as e:
            error_msg = f"処理中にエラーが発生しました:\n{str(e)}"
            self.safe_after(lambda: messagebox.showerror("エラー", error_msg))
            self.safe_after(lambda: self.reset_ui_after_task("エラーが発生しました"))

    #結果表示
    def show_result(self, text):
        self.result_area.delete("1.0", ctk.END)
        self.progress.set(1.0) 
        self.result_area.insert(ctk.END, text)
        self.transcript_text = text
        self.summary_text = ""
        try:
            auto_save_path = self.auto_save_text(text)
            status_text = f"完了しました / 自動保存: {os.path.basename(auto_save_path)}" if auto_save_path else "完了しました"
        except Exception as e:
            auto_save_path = None
            status_text = "完了しました / 自動保存に失敗しました"
            messagebox.showwarning("自動保存エラー", f"文字起こしは完了しましたが、自動保存に失敗しました:\n{str(e)}")

        # ボタン群を有効化
        self.status_label.configure(text=status_text)
        self.select_btn.configure(state="normal")
        self.run_btn.configure(state="normal")
        self.summary_prompt_menu.configure(state="normal")
        self.summary_mode_menu.configure(state="normal")
        self.update_action_buttons()
        self.cancel_btn.configure(state="disabled")
        self.active_task = None
        if auto_save_path:
            messagebox.showinfo("成功", f"文字起こしが完了しました。\n自動保存しました:\n{os.path.basename(auto_save_path)}")
        else:
            messagebox.showinfo("成功", "文字起こしが完了しました")

    #要約処理の実行
    def start_summarize_thread(self): 
        text_to_summarize = self.get_text_for_summary()
        if not text_to_summarize:
            messagebox.showwarning("警告", "要約するテキストがありません。先に文字起こしを完了してください")
            return

        # アプリ起動後にJSONを書き換えた場合も、次の要約実行でcontext_lengthを反映する
        self.load_runtime_config()
        self.save_config()
        self.cancel_event = threading.Event()
        self.active_task = "summarize"
        
        #要約処理中はボタンを非アクティブ化
        self.select_btn.configure(state="disabled")
        self.run_btn.configure(state="disabled") 
        self.summarize_btn.configure(state="disabled") 
        self.summary_prompt_menu.configure(state="disabled")
        self.summary_mode_menu.configure(state="disabled")
        self.save_txt_btn.configure(state="disabled") 
        self.save_word_btn.configure(state="disabled") 
        self.extract_speaker_btn.configure(state="disabled")
        self.apply_speaker_btn.configure(state="disabled")
        self.cancel_btn.configure(state="normal")
        self.status_label.configure(text="LM Studioで要約中...") 
        self.progress.configure(mode="determinate")
        self.progress.set(0)

        selected_prompt_key = self.summary_prompt_var.get()
        system_prompt = self.summary_prompts.get(selected_prompt_key, self.summary_prompts["標準要約（事実のみ簡潔に）"])
        summary_mode = self.summary_mode_var.get()

        self.processing_thread = threading.Thread(
            target=self.run_summarize_process,
            args=(text_to_summarize, system_prompt, summary_mode),
        )
        self.processing_thread.daemon = True
        self.processing_thread.start()  

    #LM Studioと連携した要約処理の実行
    def run_summarize_process(self, text, system_prompt, summary_mode):
        try:
            self.ensure_not_cancelled()
            client = openai.OpenAI(
                base_url="http://localhost:1234/v1",
                api_key="lm-studio",
                timeout=self.summary_timeout_seconds,
                max_retries=0,
            )

            def request_summary(messages, max_tokens=None):
                request_params = {
                    "model": "local-model",
                    "messages": messages,
                    "temperature": 0.3,
                }
                if max_tokens:
                    request_params["max_tokens"] = max_tokens

                response = client.chat.completions.create(
                    **request_params
                )
                return response.choices[0].message.content.strip()

            # 分割要約を階層的に再分割して統合
            def make_summary_groups(summary_items, max_chars=None, max_items=None):
                max_chars = max_chars or self.summary_merge_chunk_chars
                max_items = max_items or self.summary_merge_group_items
                groups = []
                current_group = []
                current_length = 0

                for summary_index, summary in enumerate(summary_items, start=1):
                    parts = self.split_text_for_summary(summary.strip(), max_chars=max_chars)
                    for part_index, part in enumerate(parts, start=1):
                        label = f"要約 {summary_index}"
                        if len(parts) > 1:
                            label = f"{label}-{part_index}"
                        item = f"【{label}】\n{part}"
                        item_length = len(item) + 2

                        should_start_next_group = (
                            current_group and
                            (current_length + item_length > max_chars or len(current_group) >= max_items)
                        )
                        if should_start_next_group:
                            groups.append(current_group)
                            current_group = []
                            current_length = 0

                        current_group.append(item)
                        current_length += item_length

                if current_group:
                    groups.append(current_group)

                return groups

            def compact_summary_if_needed(summary, target_chars, max_tokens, level, group_index):
                summary = summary.strip()
                if len(summary) <= target_chars:
                    return summary

                self.ensure_not_cancelled()
                self.safe_after(
                    lambda level=level, group_index=group_index:
                    self.status_label.configure(
                        text=f"長い要約を圧縮中... 第{level}階層 ({group_index})"
                    )
                )
                return request_summary([
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": (
                        f"以下の要約は長すぎます。重要な事実・決定事項・論点を残し、"
                        f"重複と細部を削って、必ず{target_chars}字以内に圧縮してください。\n\n"
                        f"{summary}"
                    )}
                ], max_tokens=max_tokens)

            def merge_summaries_hierarchically(summary_items):
                current_summaries = [summary.strip() for summary in summary_items if summary.strip()]
                if not current_summaries:
                    raise ValueError("統合する要約がありません。")

                level = 1
                while len(current_summaries) > 1:
                    if level > self.summary_max_merge_levels:
                        raise RuntimeError(
                            "中間要約が長すぎるため、階層統合が収束しませんでした。"
                            "分割サイズや中間要約の目安文字数を小さくして再実行してください。"
                        )

                    self.ensure_not_cancelled()
                    groups = make_summary_groups(current_summaries)
                    next_summaries = []
                    is_final_level = len(groups) == 1

                    for group_index, group in enumerate(groups, start=1):
                        self.ensure_not_cancelled()
                        self.safe_after(
                            lambda level=level, group_index=group_index, total=len(groups):
                            self.status_label.configure(
                                text=f"LM Studioで階層統合中... 第{level}階層 ({group_index}/{total})"
                            )
                        )
                        merge_progress = min(0.95, 0.72 + (level - 1) * 0.06 + (group_index / len(groups)) * 0.05)
                        self.safe_after(lambda merge_progress=merge_progress: self.progress.set(merge_progress))

                        grouped_text = "\n\n".join(group)
                        if is_final_level:
                            target_chars = self.summary_final_target_chars
                            max_tokens = self.summary_final_max_tokens
                            instruction = (
                                "以下は長い文字起こしを段階的に要約したものです。"
                                "重複を整理し、重要な事実・決定事項・論点を落とさず、"
                                f"全体として一貫した最終要約に統合してください。"
                                f"出力は{target_chars}字以内を目安にしてください。\n\n"
                            )
                        else:
                            target_chars = self.summary_intermediate_target_chars
                            max_tokens = self.summary_intermediate_max_tokens
                            instruction = (
                                "以下は長い文字起こしを分割要約したものの一部です。"
                                "次の統合段階で扱いやすい中間要約として、重複を整理し、"
                                f"重要な事実・決定事項・論点を落とさず統合してください。"
                                f"出力は必ず{target_chars}字以内にしてください。\n\n"
                            )

                        merged_summary = request_summary([
                            {"role": "system", "content": system_prompt},
                            {"role": "user", "content": f"{instruction}{grouped_text}"}
                        ], max_tokens=max_tokens)
                        merged_summary = compact_summary_if_needed(
                            merged_summary,
                            target_chars,
                            max_tokens,
                            level,
                            group_index,
                        )
                        next_summaries.append(merged_summary)

                    current_summaries = next_summaries
                    level += 1

                return current_summaries[0]

            # 一括要約
            if summary_mode == "一括要約（短文向け）":
                self.safe_after(lambda: self.status_label.configure(text="LM Studioで一括要約中..."))
                self.safe_after(lambda: self.progress.set(0.2))
                summary_result = request_summary([
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": (
                        "以下の文字起こしテキストを、指定された方針に沿って要約してください。\n\n"
                        f"{text}"
                    )}
                ], max_tokens=self.summary_final_max_tokens)
                self.ensure_not_cancelled()
                self.safe_after(lambda: self.progress.set(1.0))
                self.safe_after(lambda: self.show_summary_result(summary_result))
                return
            
            # 分割要約処理
            chunks = self.split_text_for_summary(text)
            if not chunks:
                raise ValueError("要約するテキストがありません。")

            summaries = []
            initial_progress_limit = 0.70

            for index, chunk in enumerate(chunks, start=1):
                self.ensure_not_cancelled()
                self.safe_after(lambda index=index, total=len(chunks): self.status_label.configure(
                    text=f"LM Studioで分割要約中... ({index}/{total})"
                ))
                self.safe_after(lambda index=index: self.progress.set(((index - 1) / len(chunks)) * initial_progress_limit))

                chunk_prompt = (
                    f"以下は長い文字起こしの一部です。全体の一部であることを前提に、"
                    f"重要な事実・決定事項・論点を落とさず要約してください。"
                    f"この段階の出力は必ず{self.summary_intermediate_target_chars}字以内にしてください。\n\n"
                    f"【分割 {index}/{len(chunks)}】\n\n{chunk}"
                )
                summaries.append(request_summary([
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": chunk_prompt}
                ], max_tokens=self.summary_intermediate_max_tokens))

            self.ensure_not_cancelled()

            if len(summaries) == 1:
                summary_result = summaries[0]
            else:
                self.safe_after(lambda: self.status_label.configure(text="LM Studioで統合要約中..."))
                self.safe_after(lambda: self.progress.set(initial_progress_limit))
                summary_result = merge_summaries_hierarchically(summaries)

            self.safe_after(lambda: self.progress.set(1.0))

            self.ensure_not_cancelled()
            self.safe_after(lambda: self.show_summary_result(summary_result))

        except CancelledError:
            self.safe_after(lambda: self.reset_ui_after_task("要約処理が中止されました"))

        except openai.APITimeoutError:
            error_msg = (
                f"LM Studio への要約リクエストがタイムアウトしました。\n"
                f"{self.summary_timeout_seconds}秒以内に応答がありませんでした。"
            )
            self.safe_after(lambda: messagebox.showerror("タイムアウト", error_msg))
            self.safe_after(lambda: self.reset_ui_after_task("要約がタイムアウトしました"))


        except Exception as e:
            error_msg = f"要約処理中にエラーが発生しました:\n{str(e)}"
            self.safe_after(lambda: messagebox.showerror("エラー", error_msg))
            self.safe_after(lambda: self.reset_ui_after_task("要約に失敗しました"))

    #要約結果をフォーマット
    def show_summary_result(self, summary): 
        self.summary_text = summary.strip()
        transcript_text = self.strip_summary_section(self.result_area.get("1.0", ctk.END))
        if not transcript_text:
            transcript_text = self.transcript_text.strip()
        self.transcript_text = transcript_text

        display_text = transcript_text + self.format_summary_section(self.summary_text)
        self.result_area.delete("1.0", ctk.END)
        self.result_area.insert(ctk.END, display_text)
        self.result_area.see(ctk.END) #スクロール

        self.reset_ui_after_task("要約が完了しました", keep_status=True)
        self.status_label.configure(text="要約が完了しました")
        messagebox.showinfo("成功", "要約が完了しました。")

    #保存処理
    def save_text(self): 
        text = self.result_area.get("1.0", ctk.END).strip()
        if not text:
            messagebox.showwarning("警告", "保存するテキストがありません。")
            return
        
        path = filedialog.asksaveasfilename(defaultextension=".txt", filetypes=[("Text Files", "*.txt")], title="テキスト形式で保存")
        if path:
            try:
                with open(path, "w", encoding="utf-8") as f:
                    f.write(text)
                messagebox.showinfo("成功", f"テキストファイルを保存しました:\n{os.path.basename(path)}")
            except Exception as e:
                messagebox.showerror("エラー", f"テキストファイルの保存に失敗しました:\n{str(e)}")

    def save_word(self): 
        text_to_save = self.result_area.get("1.0", ctk.END).strip()
        if not text_to_save:
            messagebox.showwarning("警告", "保存するテキストがありません。")
            return
        path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Document", "*.docx")], title="Word形式で保存")
        if path:
            try:
                doc = docx.Document()
                doc.add_heading("文字起こし・話者分離 Pro", 0)
                for line in text_to_save.split("\n"):
                    doc.add_paragraph(line)
                doc.save(path)
                messagebox.showinfo("成功", f"Wordファイルを保存しました:\n{os.path.basename(path)}")
            except Exception as e:
                messagebox.showerror("エラー", f"Wordファイルの保存に失敗しました:\n{str(e)}")

    # 辞書の読込処理
    def load_dictionary(self):
        path = filedialog.askopenfilename(defaultextension=".txt", filetypes=[("Text Files", "*.txt"), ("CSV Files", "*.csv"), ("All files", "*.*")], title="辞書ファイルを読込")
        if path:
            try:
                with open(path, "r", encoding="utf-8") as f:
                    content = f.read().strip()
                if content:
                    replacements, _ = self.parse_dictionary_text(
                        content,
                        allow_csv_pairs=path.lower().endswith(".csv"),
                    )
                    normalized_content = self.format_dictionary_text(replacements)
                    self.prompt_entry.delete(0, ctk.END)
                    self.prompt_entry.insert(0, normalized_content)
                    messagebox.showinfo("成功", f"辞書を読み込みました:\n{os.path.basename(path)}")
                else:
                    messagebox.showinfo("確認", "選択したファイルは空でした。")
            except Exception as e:
                messagebox.showerror("エラー", f"辞書の読込に失敗しました:\n{str(e)}")

    # 辞書の保存処理
    def save_dictionary(self):
        text = self.prompt_entry.get().strip()
        if not text:
            messagebox.showwarning("警告", "保存する辞書データがありません。")
            return
        
        path = filedialog.asksaveasfilename(defaultextension=".txt", filetypes=[("Text Files", "*.txt"), ("CSV Files", "*.csv")], title="辞書データを保存")
        if path:
            try:
                replacements, _ = self.parse_dictionary_text(text)
                normalized_text = self.format_dictionary_text(replacements, separator="\n")
                with open(path, "w", encoding="utf-8") as f:
                    f.write(normalized_text)
                messagebox.showinfo("成功", f"辞書ファイルを保存しました:\n{os.path.basename(path)}")
            except Exception as e:
                messagebox.showerror("エラー", f"辞書の保存に失敗しました:\n{str(e)}")

if __name__ == "__main__":
    root = ctk.CTk()
    app = Whisperapp(root)
    root.mainloop()
