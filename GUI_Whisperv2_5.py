#文字おこし、話者分離、要約に対応したバージョン（要約はLM Studio使用）
#使用時はLM Studioでローカルサーバーを起動し、llamaを使用すること
#Windowsでは動作しないので注意
import customtkinter as ctk
from tkinter import filedialog, messagebox
import mlx_whisper
import threading
import os
import docx
import torch
import gc
import openai
import re
import ctypes

class CancelledError(BaseException):
    pass

# --- 環境設定 ---
os.environ["HF_HUB_OFFLINE"] = "0"
from pyannote.audio import Pipeline

# --- デザインの基本設定 ---
ctk.set_appearance_mode("light") 
ctk.set_default_color_theme("blue")

class Whisperapp:
    def __init__(self, root):
        self.root = root
        self.root.title("MLX Whisper Pro - 文字起こし & 要約")
        # ★ 縦幅を1000から800に変更し、MacBookの画面に収まるようにしました
        self.root.geometry("1000x800")
        self.root.configure(fg_color="#FFFFFF")

        # Mac用フォント設定
        self.font_title = ("Hiragino Sans", 15, "bold")
        self.font_main = ("Hiragino Sans", 13)
        self.font_text = ("Menlo", 13)

        self.diarization_pipeline = None

        self.models = {
            "Turbo（高速・高精度）": "mlx-community/whisper-large-v3-turbo",
            "Large v3（最高精度・低速）": "mlx-community/whisper-large-v3-mlx-4bit",
            "Small（バランス）": "mlx-community/whisper-small-mlx",
            "Base（軽量・高速）": "mlx-community/whisper-base-mlx",
            "Tiny（最速・低精度）": "mlx-community/whisper-tiny-mlx"
        }

        self.summary_prompts = {
            "標準要約（事実のみ簡潔に）": "あなたは優秀なアシスタントです。以下の文字起こしテキストを要約してください。不要な相槌や重複した議論は削り、事実のみを正確にまとめてください。",
            "議事録（箇条書き）": "あなたは優秀な書記です。以下の文字起こしテキストから重要なポイントを抽出し、見出しと箇条書きを使った分かりやすい議事録を作成してください。",
            "決定事項・ToDoの抽出": "以下の文字起こしテキストから、「決定された事項」と「各担当者の次のアクション(ToDo)」のみを箇条書きで分かりやすく抽出してください。",
            "行政報告書形式（きわめてフォーマル）": "あなたは優秀な行政官です。以下の文字起こしテキストを、公式な行政報告書にふさわしい、厳格で正確な公用文にて要約してください。客観的な事実関係を整理し、感情表現や重複を完全に排除して簡潔にまとめてください。"
        }

        # --- メインレイアウト ---
        self.main_frame = ctk.CTkFrame(root, fg_color="transparent")
        self.main_frame.pack(fill=ctk.BOTH, expand=True, padx=30, pady=15) # 余白を少し圧縮

        # 1. ファイル選択セクション
        self.label_step1 = ctk.CTkLabel(self.main_frame, text="1. 音声・動画ファイルを選択", font=self.font_title, text_color="#1D1D1F")
        self.label_step1.pack(anchor="w", pady=(0, 5))

        self.select_btn = ctk.CTkButton(self.main_frame, text="ファイルを選択", command=self.select_file, 
                                        font=self.font_main, height=35, corner_radius=8, 
                                        fg_color="#F5F5F7", text_color="#007AFF", hover_color="#E5E5E7")
        self.select_btn.pack(fill=ctk.X)

        self.file_path_label = ctk.CTkLabel(self.main_frame, text="ファイルが選択されていません", 
                                            text_color="#8E8E93", font=("Hiragino Sans", 12))
        self.file_path_label.pack(pady=(2, 10))

        # 2. 設定エリア
        self.config_frame = ctk.CTkFrame(self.main_frame, fg_color="#FBFBFD", corner_radius=15, border_width=1, border_color="#D2D2D7")
        self.config_frame.pack(fill=ctk.X, pady=5, ipady=5)

        # Hugging Face トークン
        ctk.CTkLabel(self.config_frame, text="Hugging Face トークン", font=self.font_main).grid(row=0, column=0, padx=15, pady=(10, 5), sticky="w")
        self.token_entry = ctk.CTkEntry(self.config_frame, width=350, placeholder_text="Token (optional)",
                                         corner_radius=8, border_color="#D2D2D7", fg_color="#FFFFFF", show="*")
        self.token_entry.grid(row=0, column=1, padx=15, pady=(10, 5), sticky="w")

        # モデル選択
        ctk.CTkLabel(self.config_frame, text="使用するAIモデル", font=self.font_main).grid(row=1, column=0, padx=15, pady=5, sticky="w")
        self.model_var = ctk.StringVar(value="Turbo（高速・高精度）")
        self.model_menu = ctk.CTkComboBox(self.config_frame, values=list(self.models.keys()), 
                                          variable=self.model_var, width=350, corner_radius=8)
        self.model_menu.grid(row=1, column=1, padx=15, pady=5, sticky="w")

        # 話者分離オプション
        self.diarize_var = ctk.BooleanVar(value=False)
        self.diarize_check = ctk.CTkCheckBox(self.config_frame, text="話者分離を有効にする", variable=self.diarize_var, font=self.font_main)
        self.diarize_check.grid(row=2, column=0, padx=15, pady=(5, 10), sticky="w")

        self.num_speakers_var = ctk.StringVar(value="自動")
        self.num_speakers_menu = ctk.CTkComboBox(self.config_frame, values=["自動", "2", "3", "4", "5", "6", "7", "8"], variable=self.num_speakers_var, width=100)
        self.num_speakers_menu.grid(row=2, column=1, padx=15, pady=(5, 10), sticky="w")

        # 3. 専門用語辞書
        self.dict_frame = ctk.CTkFrame(self.main_frame, fg_color="transparent")
        self.dict_frame.pack(fill=ctk.X, pady=(10, 5))

        ctk.CTkLabel(self.dict_frame, text="3. 専門用語・略語の登録（任意）", 
                     font=self.font_title, text_color="#1D1D1F").pack(side=ctk.LEFT)

        self.save_dict_btn = ctk.CTkButton(self.dict_frame, text="保存", command=self.save_dictionary, width=60, height=28, font=self.font_main, fg_color="#F5F5F7", text_color="#007AFF", hover_color="#E5E5E7")
        self.save_dict_btn.pack(side=ctk.RIGHT, padx=(5, 0))

        self.load_dict_btn = ctk.CTkButton(self.dict_frame, text="読込", command=self.load_dictionary, width=60, height=28, font=self.font_main, fg_color="#F5F5F7", text_color="#007AFF", hover_color="#E5E5E7")
        self.load_dict_btn.pack(side=ctk.RIGHT)

        self.prompt_entry = ctk.CTkEntry(self.main_frame, placeholder_text="例）AI: 人工知能, ML: 機械学習", 
                                         corner_radius=8, height=35, border_color="#D2D2D7", fg_color="#FFFFFF")
        self.prompt_entry.pack(fill=ctk.X, pady=(0, 10))

        # 実行・キャンセルボタン群
        self.run_frame = ctk.CTkFrame(self.main_frame, fg_color="transparent")
        self.run_frame.pack(fill=ctk.X, pady=5)

        self.run_btn = ctk.CTkButton(self.run_frame, text="文字起こしを開始する", command=self.start_thread, state="disabled",
                                     height=45, corner_radius=22, font=("Hiragino Sans", 15, "bold"), fg_color="#007AFF", hover_color="#005BB5")
        self.run_btn.pack(side=ctk.LEFT, expand=True, fill=ctk.X, padx=(0, 5))

        self.cancel_btn = ctk.CTkButton(self.run_frame, text="中止", command=self.cancel_process, state="disabled", width=80,
                                     height=45, corner_radius=22, font=("Hiragino Sans", 15, "bold"), text_color="#FFFFFF", fg_color="#FF3B30", hover_color="#D70015")
        self.cancel_btn.pack(side=ctk.LEFT, padx=(5, 0))

        # 結果表示エリア
        self.result_area = ctk.CTkTextbox(self.main_frame, height=150, corner_radius=10, border_width=1, 
                                          border_color="#D2D2D7", font=self.font_text, fg_color="#FFFFFF", text_color="#1D1D1F")
        self.result_area.pack(fill=ctk.BOTH, expand=True, pady=10)

        # 保存・要約ボタン群
        self.action_frame = ctk.CTkFrame(self.main_frame, fg_color="transparent")
        self.action_frame.pack(fill=ctk.X, pady=(0, 5))

        self.save_txt_btn = ctk.CTkButton(self.action_frame, text="Text保存", command=self.save_text, state="disabled", corner_radius=8, fg_color="#8E8E93")
        self.save_txt_btn.pack(side=ctk.LEFT, padx=5, expand=True, fill=ctk.X)

        self.save_word_btn = ctk.CTkButton(self.action_frame, text="Word保存", command=self.save_word, state="disabled", corner_radius=8, fg_color="#8E8E93")
        self.save_word_btn.pack(side=ctk.LEFT, padx=5, expand=True, fill=ctk.X)

        # 要約プロンプトの選択
        self.summary_prompt_var = ctk.StringVar(value="標準要約（事実のみ簡潔に）")
        self.summary_prompt_menu = ctk.CTkComboBox(self.action_frame, values=list(self.summary_prompts.keys()), 
                                                   variable=self.summary_prompt_var, corner_radius=8, width=220)
        self.summary_prompt_menu.pack(side=ctk.LEFT, padx=(15, 5))

        self.summarize_btn = ctk.CTkButton(self.action_frame, text="LM Studioで要約", command=self.start_summarize_thread, state="disabled", corner_radius=8, fg_color="#34C759", hover_color="#248A3D")
        self.summarize_btn.pack(side=ctk.LEFT, padx=(0, 5), expand=True, fill=ctk.X)

        # ステータスバー（一番下に固定）
        self.status_label = ctk.CTkLabel(root, text="準備完了", fg_color="#F5F5F7", height=25, font=("Hiragino Sans", 11))
        self.status_label.pack(side=ctk.BOTTOM, fill=ctk.X)
        self.progress = ctk.CTkProgressBar(root, height=4, corner_radius=0, fg_color="#E5E5E7", progress_color="#007AFF")
        self.progress.pack(side=ctk.BOTTOM, fill=ctk.X)
        self.progress.set(0)

        self.filepath = ""

    #実行ファイルフェーズ
    def select_file(self):
        file_types = [("Audio/Video files", "*.mp3 *.wav *.flac *.m4a *.mp4 *.mov *.mkv"), ("All files", "*.*")]
        selected_path = filedialog.askopenfilename(filetypes=file_types)
        if selected_path:
            self.filepath = selected_path
            self.file_path_label.configure(text=os.path.basename(self.filepath), text_color="black")
            self.run_btn.configure(state="normal")

    #処理中はボタンを非アクティブ化
    def start_thread(self):
        self.run_btn.configure(state="disabled")
        self.save_txt_btn.configure(state="disabled") 
        self.save_word_btn.configure(state="disabled")
        self.summary_prompt_menu.configure(state="disabled")
        self.summarize_btn.configure(state="disabled")
        self.cancel_btn.configure(state="normal")
        self.result_area.delete("1.0", ctk.END)
        self.status_label.configure(text="処理しています…")
        self.progress.set(0)

        self.processing_thread = threading.Thread(target=self.run_process)
        self.processing_thread.daemon = True 
        self.processing_thread.start()

    def cancel_process(self):
        if hasattr(self, 'processing_thread') and self.processing_thread and self.processing_thread.is_alive():
            self.status_label.configure(text="中止しています...")
            self.cancel_btn.configure(state="disabled")
            thread_id = self.processing_thread.ident
            res = ctypes.pythonapi.PyThreadState_SetAsyncExc(ctypes.c_long(thread_id), ctypes.py_object(CancelledError))
            if res > 1:
                ctypes.pythonapi.PyThreadState_SetAsyncExc(ctypes.c_long(thread_id), 0)

    def run_process(self):
        try:
            if not self.filepath or not os.path.isfile(self.filepath):
                raise ValueError("ファイルが選択されていないか、存在しません。もう一度ファイルを選択し直してください。")

            token = self.token_entry.get().strip()
            do_diarize = self.diarize_var.get()
            user_prompt = self.prompt_entry.get().strip()
            
            prompt_dict = {} 
            hint_words = [] 

            if user_prompt:
                for item in user_prompt.split(","):
                    if ":" in item:
                        wrong, correct = item.split(":", 1)
                        prompt_dict[wrong.strip()] = correct.strip()
                        hint_words.append(wrong.strip()) 
                    else:
                        word = item.strip()
                        if word:
                            prompt_dict[word] = word
                            hint_words.append(word)

                initial_prompt_str = ", ".join(hint_words) if hint_words else None
                prompt_for_whisper = {f"replace:{k}": v for k, v in prompt_dict.items()} if prompt_dict else None
            else:
                initial_prompt_str = None
                prompt_for_whisper = None

            # 1.Whisperで文字起こし
            self.root.after(0, lambda: self.status_label.configure(text="Whisperで文字起こし中... しばらくお待ちください..."))
            model_path = self.models[self.model_var.get()]

            #長時間データの場合の安定化オプション
            whisper_options = {
                "language": "ja",                       #言語を日本語に固定
                "condition_on_previous_text": False,   #前の文脈を引きずらない
                "compression_ratio_threshold": 2.0,     #読み間違いのループを検知してリセット
                "no_speech_threshold": 0.4,             #無音区間での暴走を防ぐ
                "logprob_threshold": -0.5,              #無音区間の暴走抑止
                "temperature": 0.0,                      #ランダムな記号の生成抑止
                "word_timestamps": False,               #タイムスタンプ計算のバグによる記号生成の抑止
                "best_of": 5                            #5つの候補から最も適切なものを選択
            }
            
            if initial_prompt_str: 
                whisper_result = mlx_whisper.transcribe(
                    self.filepath, path_or_hf_repo=model_path, initial_prompt=initial_prompt_str, **whisper_options)
            else:
                whisper_result = mlx_whisper.transcribe(
                    self.filepath, path_or_hf_repo=model_path, initial_prompt="こんにちは。", **whisper_options)

            if prompt_for_whisper: 
                for segment in whisper_result["segments"]:
                    text = segment["text"]
                    for wrong, correct in prompt_dict.items():
                        text = text.replace(wrong, correct) 
                    segment["text"] = text

                full_text = whisper_result["text"]
                for wrong, correct in prompt_dict.items():
                    full_text = full_text.replace(wrong, correct)
                whisper_result["text"] = full_text

            final_text = ""
            gc.collect() 
            if torch.backends.mps.is_available():
                torch.mps.empty_cache() 

            # 2.話者分離の実行
            if do_diarize:
                self.root.after(0, lambda: self.status_label.configure(text="話者解析中... (初回はモデルをロードします)"))
                if self.diarization_pipeline is None:
                    auth_param = token if token else True 
                    self.diarization_pipeline = Pipeline.from_pretrained("pyannote/speaker-diarization-3.1", token=auth_param)

                if torch.backends.mps.is_available():
                    self.diarization_pipeline.to(torch.device("mps"))

                def diarization_hook(*args, **kwargs):
                    completed = kwargs.get("completed") or 0
                    total = kwargs.get("total") or 1

                    if total >0:
                        # CTkProgressBarは 0.0〜1.0 の値を設定
                        percentage_float = completed / total
                        percentage_int = int(percentage_float * 100)

                        self.root.after(0, lambda p=percentage_float: self.progress.set(p))
                        self.root.after(0, lambda p=percentage_int: self.status_label.configure(
                            text=f"話者解析中...しばらくお待ちください... ({p}%)"))

                batch_size = 16 
                num_spk = None if self.num_speakers_var.get() == "自動" else int(self.num_speakers_var.get())
                
                diarization_result = self.diarization_pipeline(self.filepath, num_speakers=num_spk, batch_size=batch_size, hook=diarization_hook)

                tracks = [] 
                if hasattr(diarization_result, "itertracks"): 
                    for turn, _, spk_label in diarization_result.itertracks(yield_label=True):
                        tracks.append((turn, spk_label))
                elif hasattr(diarization_result, "speaker_diarization"): 
                    data = diarization_result.speaker_diarization
                    if hasattr(data, "itertracks"):
                        for turn, _, spk_label in data.itertracks(yield_label=True):
                            tracks.append((turn, spk_label))
                    else:
                        for turn, spk_label in data:
                            tracks.append((turn, spk_label))
                else:
                    raise ValueError("Pyannoteの結果の形式が不明です。")

                for segment in whisper_result["segments"]:
                    start_time = segment["start"]
                    end_time = segment["end"]
                    text = segment["text"].strip()
                    speaker = "Unknown"
                    max_overlap = -1.0

                    for turn, spk_label in tracks:
                        overlap_start = max(start_time, turn.start) 
                        overlap_end = min(end_time, turn.end)
                        overlap = overlap_end - overlap_start
                        if overlap > max_overlap and overlap > 0: 
                            max_overlap = overlap
                            speaker = spk_label

                    time_str = f"[{int(start_time//60):02d}:{int(start_time%60):02d}]"
                    final_text += f"【{speaker}】{time_str} {text}\n"
            else:
                final_text = whisper_result["text"].strip()

            final_text = re.sub(r'([^\s\n])\1{9,}', r'\1', final_text) #同じ文字が9回以上続いたら1文字に圧縮
            final_text = re.sub(r'([;:.])\1{5,}', r'\1', final_text) #特定の記号も3つにまとめる
                
            self.root.after(0, lambda: self.show_result(final_text))

        except CancelledError:
            gc.collect()
            if torch.backends.mps.is_available(): torch.mps.empty_cache()
            self.root.after(0, lambda: self.status_label.configure(text="処理が中止されました"))
            self.root.after(0, lambda: self.progress.set(0))
            self.root.after(0, lambda: self.run_btn.configure(state="normal"))
            self.root.after(0, lambda: self.cancel_btn.configure(state="disabled"))

        #エラー表示
        except Exception as e:
            error_msg = str(e)
            self.root.after(0, lambda: self.progress.set(0)) 
            self.root.after(0, lambda msg=error_msg: messagebox.showerror("エラー", f"処理中にエラーが発生しました:\n{msg}"))
            self.root.after(0, lambda: self.status_label.configure(text="エラーが発生しました"))
            self.root.after(0, lambda: self.run_btn.configure(state="normal"))
            self.root.after(0, lambda: self.cancel_btn.configure(state="disabled"))

    #結果表示
    def show_result(self, text): 
        self.progress.set(1.0) 
        self.result_area.insert(ctk.END, text)
        self.status_label.configure(text="完了しました")
        self.run_btn.configure(state="normal")
        self.save_txt_btn.configure(state="normal") 
        self.save_word_btn.configure(state="normal") 
        self.summary_prompt_menu.configure(state="normal")
        self.summarize_btn.configure(state="normal") 
        self.cancel_btn.configure(state="disabled")
        messagebox.showinfo("成功", "文字起こしが完了しました")

    #要約処理の実行
    def start_summarize_thread(self): 
        text_to_summarize = self.result_area.get("1.0", ctk.END).strip()
        if not text_to_summarize:
            messagebox.showwarning("警告", "要約するテキストがありません。")
            return
        
        #要約処理中はボタンを非アクティブ化
        self.run_btn.configure(state="disabled") 
        self.summarize_btn.configure(state="disabled") 
        self.summary_prompt_menu.configure(state="disabled")
        self.save_txt_btn.configure(state="disabled") 
        self.save_word_btn.configure(state="disabled") 
        self.cancel_btn.configure(state="normal")
        self.status_label.configure(text="LM Studioで要約中...") 
        self.progress.configure(mode="indeterminate")
        self.progress.start() 

        self.processing_thread = threading.Thread(target=self.run_summarize_process, args=(text_to_summarize,))
        self.processing_thread.daemon = True
        self.processing_thread.start()  

    #LM Studioと連携した要約処理の実行
    def run_summarize_process(self, text):
        try:
            client = openai.OpenAI(base_url="http://localhost:1234/v1", api_key="lm-studio") 
            
            # 要約プロンプトをUIから取得
            selected_prompt_key = self.summary_prompt_var.get()
            system_prompt = self.summary_prompts.get(selected_prompt_key, self.summary_prompts["標準要約（事実のみ簡潔に）"])

            response = client.chat.completions.create(
                model="local-model", 
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": f"以下の文字起こしテキストを要約してください:\n\n{text}"}
                ],
                temperature=0.3 #0.0~1.0の間で調整。低い程、事実に基づく堅い要約になる
            )

            summary_result = response.choices[0].message.content 
            self.root.after(0, lambda: self.show_summary_result(summary_result)) 

        except CancelledError:
            self.root.after(0, lambda: self.status_label.configure(text="要約処理が中止されました"))

        except Exception as e:
            error_msg = f"要約処理中にエラーが発生しました:\n{str(e)}"
            self.root.after(0, lambda msg=error_msg: messagebox.showerror("エラー", msg))

        finally:
            self.root.after(0, lambda: self.progress.stop()) 
            self.root.after(0, lambda: self.progress.configure(mode="determinate"))
            self.root.after(0, lambda: self.progress.set(0))
            self.root.after(0, lambda: self.status_label.configure(text="準備完了")) 
            self.root.after(0, lambda: self.run_btn.configure(state="normal")) 
            self.root.after(0, lambda: self.summarize_btn.configure(state="normal")) 
            self.root.after(0, lambda: self.summary_prompt_menu.configure(state="normal"))
            self.root.after(0, lambda: self.save_txt_btn.configure(state="normal")) 
            self.root.after(0, lambda: self.save_word_btn.configure(state="normal")) 
            self.root.after(0, lambda: self.cancel_btn.configure(state="disabled"))

    #要約結果をフォーマット
    def show_summary_result(self, summary): 
        formatted_summary = f"\n \n=========================\n【要約結果】\n=========================\n{summary}\n\n" 
        self.result_area.insert(ctk.END, formatted_summary) 
        self.result_area.see(ctk.END) #スクロール

        self.status_label.configure(text="要約が完了しました") 
        messagebox.showinfo("成功", "要約が完了しました。")

    #保存処理
    def save_text(self): 
        text = self.result_area.get("1.0", ctk.END).strip()
        if not text:
            messagebox.showwarning("警告", "保存するテキストがありません。")
            return
        
        path = filedialog.asksaveasfilename(defaultextension=".txt", filetypes=[("Text Files", "*.txt")], title="テキスト形式で保存")
        if path:
            try:
                with open(path, "w", encoding="utf-8") as f:
                    f.write(text)
                messagebox.showinfo("成功", f"テキストファイルを保存しました:\n{os.path.basename(path)}")
            except Exception as e:
                messagebox.showerror("エラー", f"テキストファイルの保存に失敗しました:\n{str(e)}")

    def save_word(self): 
        text_to_save = self.result_area.get("1.0", ctk.END).strip()
        if not text_to_save:
            messagebox.showwarning("警告", "保存するテキストがありません。")
            return
        path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Document", "*.docx")], title="Word形式で保存")
        if path:
            try:
                doc = docx.Document()
                doc.add_heading("文字起こし・話者分離 Pro", 0)
                for line in text_to_save.split("\n"):
                    doc.add_paragraph(line)
                doc.save(path)
                messagebox.showinfo("成功", f"Wordファイルを保存しました:\n{os.path.basename(path)}")
            except Exception as e:
                messagebox.showerror("エラー", f"Wordファイルの保存に失敗しました:\n{str(e)}")

    # 辞書の読込処理
    def load_dictionary(self):
        path = filedialog.askopenfilename(defaultextension=".txt", filetypes=[("Text Files", "*.txt"), ("CSV Files", "*.csv"), ("All files", "*.*")], title="辞書ファイルを読込")
        if path:
            try:
                with open(path, "r", encoding="utf-8") as f:
                    content = f.read().strip()
                if content:
                    self.prompt_entry.delete(0, ctk.END)
                    self.prompt_entry.insert(0, content)
                    messagebox.showinfo("成功", f"辞書を読み込みました:\n{os.path.basename(path)}")
                else:
                    messagebox.showinfo("確認", "選択したファイルは空でした。")
            except Exception as e:
                messagebox.showerror("エラー", f"辞書の読込に失敗しました:\n{str(e)}")

    # 辞書の保存処理
    def save_dictionary(self):
        text = self.prompt_entry.get().strip()
        if not text:
            messagebox.showwarning("警告", "保存する辞書データがありません。")
            return
        
        path = filedialog.asksaveasfilename(defaultextension=".txt", filetypes=[("Text Files", "*.txt"), ("CSV Files", "*.csv")], title="辞書データを保存")
        if path:
            try:
                with open(path, "w", encoding="utf-8") as f:
                    f.write(text)
                messagebox.showinfo("成功", f"辞書ファイルを保存しました:\n{os.path.basename(path)}")
            except Exception as e:
                messagebox.showerror("エラー", f"辞書の保存に失敗しました:\n{str(e)}")

if __name__ == "__main__":
    root = ctk.CTk()
    app = Whisperapp(root)
    root.mainloop()
